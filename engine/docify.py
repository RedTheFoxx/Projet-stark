"""Module de génération de rapports au format docx."""

import json
import random
from datetime import datetime
from docx import Document
from docx.shared import Inches


def make_intervention_report(data: str):
    """Génère un rapport d'intervention au format docx.

    Args:
        data (str): Données d'intervention en syntaxe JSON valide qu'il faudra convertir.

    Returns:
        Document: Document Word.
    """
    converted_data = json.loads(  # On s'assure de convertir la string d'entrée en dictionnaire Python
        data
    )
    doc = Document()

    # Accès à la première section du document
    section = doc.sections[0]

    # Définition des marges en pouces
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    header_table = doc.add_table(
        rows=1, cols=3
    )  # Case 1 : la date d'intervention et case 2, le logo de stark
    header_table.cell(0, 0).text = "\n\n\nDate d'intervention :"
    header_table.cell(0, 1).text = "\n\n\n" + converted_data["Date d'intervention"]
    logo_cell = header_table.cell(0, 2)
    logo_path = "ressources/data/logo_stark.png"
    logo_paragraph = logo_cell.paragraphs[0]
    logo_run = logo_paragraph.add_run()
    logo_run.add_picture(logo_path, width=Inches(2.5))

    doc.add_heading("Rapport d'intervention", 0)
    doc.add_heading("Informations client\n", 2)

    keys_to_display = ["Libellé du site", "Ville", "Motif", "Statut"]

    table = doc.add_table(rows=len(keys_to_display), cols=2)

    for cell in table.columns[0].cells:
        cell.width = Inches(1.25)
    for cell in table.columns[1].cells:
        cell.width = Inches(6)

    for i, key in enumerate(keys_to_display):
        if key in converted_data:
            table.cell(i, 0).text = key
            table.cell(i, 1).text = converted_data[key]

    doc.add_heading("Compte-rendu d'intervention\n", 2)
    doc.add_paragraph(converted_data["Resume"])
    doc.add_paragraph(
        """Nous tenons à vous remercier pour votre confiance et restons à votre disposition pour toute information complémentaire.\nN'hésitez pas à contacter le Centre de Relation Clients (CRC) au 0 800 80 93 00, disponible 24h/24 et 7j/7.\nNous tenons également à souligner l'importance de faire suivre votre matériel par un professionnel pour garantir sa qualité et sa durabilité.\nCordialement,"""
    )

    doc.add_paragraph(converted_data["Intervenant"])

    return doc


def make_activity_report(
    data_customer_site: str, data_interventions: list, data_report: str
):
    """Génère un rapport d'activité au format docx.

    Args:
        data_customer_site (str): Le nom du client (qui est aussi le filtre principal sur les données d'entrée)
        data_interventions (list): Liste des interventions en syntaxe JSON valide = un enregistrement par élément de liste.
        data_report (str): Un bloc de texte représentant le bilan annuel généré sur la base des data_interventions

    Returns:
        Document: Document Word.
    """
    doc = Document()

    header_table = doc.add_table(
        rows=1, cols=2
    )  # Case 1 : le numéro de demande au hasard pour la démonstration et case 2, le logo de stark
    header_table.cell(0, 0).text = "N° de demande : " + str(random.randint(1, 1000000))
    logo_cell = header_table.cell(0, 1)
    logo_path = "ressources/data/logo_stark.png"
    logo_paragraph = logo_cell.paragraphs[0]
    logo_run = logo_paragraph.add_run()
    logo_run.add_picture(logo_path, width=Inches(2.5))

    doc.add_heading("Rapport d'activité", 0)
    doc.add_heading("Informations client :", 2)

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Date & heure du rapport"
    table.cell(0, 1).text = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
    table.cell(1, 0).text = "Client"
    table.cell(1, 1).text = data_customer_site

    doc.add_heading("Compte-rendu d'interventions :", 2)
    doc.add_paragraph("----------------------------------------------")

    # On trie dans l'ordre chronologique avant d'ajouter paragraphes par paragraphes
    sorted_interventions = sorted(
        [json.loads(intervention_str) for intervention_str in data_interventions],
        key=lambda x: datetime.strptime(
            x.get("Intervention du ", "01-01-1900 00:00:00"), "%d-%m-%Y %H:%M:%S"
        ),
    )

    # ---------------------------------------------------------------------------- #

    for i, intervention in enumerate(sorted_interventions):
        doc.add_heading(f"Intervention n°{i+1}", 3)
        doc.add_paragraph(
            f"Date : {intervention.get('Intervention du ', 'Non spécifiée')}"
        )
        doc.add_paragraph(f"Durée : {intervention.get('Durée', 'Non spécifiée')}")
        doc.add_paragraph(intervention.get("Résumé", "Aucun résumé disponible"))

    doc.add_heading("Bilan annuel :", 2)
    doc.add_paragraph(data_report)

    return doc


if __name__ == "__main__":

    data_customer_site = "DIEPPE CHATEAU MUSEE"

    data_interventions = [
        '{\n    "Intervention du ": "02-05-2023 09:39:00",\n    "Durée": "00:30:00",\n    "Résumé": "Le client a demandé l\'arrêt du chauffage pour permettre à l\'entreprise VALLET de remplacer les canalisations. L\'intervention a été réalisée le 2 mai 2023 et le chauffage a été coupé pour les travaux."\n}',
        '{\n    "Intervention du ": "25-01-2023 14:00:00",\n    "Durée": "00:30:00",\n    "Résumé": "Le client a signalé un problème de chauffage dans le restaurant du Forges Hôtel, mentionnant qu\'il faisait très froid. L\'intervention a été réalisée et le problème de chauffage a été résolu. Le technicien a consigné les températures."\n}',
        '{\n    "Intervention du ": "12-01-2023 11:44:00",\n    "Durée": "02:11:00",\n    "Résumé": "Le client a signalé un arrêt de chauffage malgré une intervention précédente. Nous avons ajusté le départ régulé et contrôlé la température à 18,6°C, ainsi que le bon fonctionnement général de l\'installation."\n}',
        '{\n    "Intervention du ": "21-12-2023 07:53:00",\n    "Durée": "08:45:00",\n    "Résumé": "Intervention demandée pour un bruit assourdissant dans le bureau des marchés publics à Walter Heights. Problème réglé. Intervention effectuée sans remarque ou alerte particulière de la part du technicien."\n}',
        '{\n    "Intervention du ": "29-11-2023 12:10:00",\n    "Durée": "00:00:00",\n    "Résumé": "Le client a signalé un défaut de pressostat ECS et un nombre maximum de communications par heure atteint. L\'intervention a été effectuée sans remarque ou alerte particulière de la part du technicien."\n}',
    ]

    data_report = "BLOC QUI PARLE DU RESUME ANNUEL"

    generated_doc = make_activity_report(
        data_customer_site, data_interventions, data_report
    )
    generated_doc.save("output/RA_test.docx")
